import io
import os
import docx
import pypdf
import streamlit as st
from google import genai

# 화면 설정
st.set_page_config(page_title="간호학과 포트폴리오 점검", layout="centered")
st.title("🩺 간호교육 포트폴리오 AI 점검")
st.write(
    "포트폴리오 파일들(PDF/Word)을 한꺼번에 업로드하면 AI가 통합"
    " 점검합니다."
)

# API 키 확인
api_key = st.secrets.get("GEMINI_API_KEY") or os.environ.get("GEMINI_API_KEY")
if not api_key:
  st.error(
      "API 키가 설정되지 않았습니다. Streamlit Secrets에서"
      " GEMINI_API_KEY를 설정해 주세요."
  )
  st.stop()

client = genai.Client(api_key=api_key)

# 다중 파일 업로드
uploaded_files = st.file_uploader(
    "포트폴리오 파일 선택 (.docx, .pdf) - 여러 파일 선택 가능",
    type=["docx", "pdf"],
    accept_multiple_files=True,
)


def extract_text(file):
  ext = os.path.splitext(file.name)[1].lower()
  full_text = []
  if ext == ".docx":
    doc = docx.Document(file)
    for para in doc.paragraphs:
      full_text.append(para.text)
  elif ext == ".pdf":
    reader = pypdf.PdfReader(file)
    for page in reader.pages:
      text = page.extract_text()
      if text:
        full_text.append(text)
  return "\n".join(full_text)


def create_docx(report_text):
  doc = docx.Document()
  doc.add_heading("간호교육 포트폴리오 AI 점검 리포트", level=1)
  for line in report_text.split("\n"):
    if line.startswith("# "):
      doc.add_heading(line.replace("# ", ""), level=1)
    elif line.startswith("## "):
      doc.add_heading(line.replace("## ", ""), level=2)
    elif line.startswith("### "):
      doc.add_heading(line.replace("### ", ""), level=3)
    else:
      doc.add_paragraph(line)

  bio = io.BytesIO()
  doc.save(bio)
  return bio.getvalue()


if uploaded_files:
  st.info(f"선택된 파일 수: 총 {len(uploaded_files)}개")
  for f in uploaded_files:
    st.write(f"- {f.name}")

  if st.button("AI 점검 시작"):
    with st.spinner("AI가 제출된 문서 전체를 통합 분석 중입니다..."):
      try:
        combined_text = ""
        for file in uploaded_files:
          combined_text += f"\n\n--- [파일명: {file.name}] ---\n"
          combined_text += extract_text(file)

        prompt = f"""
                당신은 간호교육인증평가 및 교육품질관리(CQI) 전문가입니다.
                제출된 문서를 바탕으로 다음 3가지 핵심 항목 및 세부 지표를 엄격히 점검해 주세요.

                [점검 항목 및 세부 평가 기준]

                1. [강의계획 및 PO-CO 연계성]
                   - 교과목 학습성과(CO)와 프로그램 학습성과(PO)의 매핑 및 연계성이 타당한가?
                   - 주차별 수업 내용 및 교수학습방법이 학습성과 달성에 적절한가?

                2. [학생 평가 도구 및 루브릭]
                   - 평가 도구, 채점 기준표(루브릭) 및 성적대별 구비 여부가 적절한가?

                3. [CQI 분석 및 미성취자 환류 관리]
                   - [학습성과별 분석]: 각 학습성과(PO/CO)별 달성도 및 성취 수준 분석이 면밀한가?
                   - [[하] 성취수준 기준 적용 적절성]: 
                     * 기본 절단값: 2023~2025 교육과정 적용 대상은 [달성도 70% 미만], 2026 교육과정 적용 대상은 [달성도 60% 미만]을 적용
                     * [휴복학생/재수강생 예외 판정 기준]: 단순 학번(학번 앞자리)이 아닌 **실제 수강 중인 학년 및 적용 교육과정**을 최우선 기준으로 인정함
                       (예: 2025학번이라도 휴학 후 2026학년도 1학년으로 복학하여 2026 교육과정 수강 시, 2026학번 기준인 60% 미만 적용을 '적합'으로 판정)
                     * 단, 학번과 적용 기준이 다를 경우 보고서 내에 복학/재수강 등 예외 사유나 학년 기준 적용 표기가 되어 있는지 확인
                   - [미성취자 관리 및 환류]: 해당 학년 기준 [하] 성취수준 학생에 대한 구체적인 지도, 재평가, 보완 프로그램 등 미성취자 관리가 철저히 수행되었는가?
                   - [차기 학기 환류계획]: 이전 CQI 반영 여부 및 차기 학기 개선계획이 실현 가능한가?

                [작성 형식]
                - 각 항목별로 [적합], [보완 필요], [미흡] 판정을 명시해 주세요.
                - 학번과 실제 수강 학년 기준 오적용 여부 및 미성취자 관리 방안 미비 시 구체적인 보완 요구사항을 명확히 제시해 주세요.

                제출 문서 내용:
                {combined_text}
                """
        response = client.models.generate_content(
            model="gemini-3.6-flash", contents=prompt
        )

        st.session_state["report"] = response.text
        st.success("점검 완료!")

      except Exception as e:
        st.error(f"분석 중 오류가 발생했습니다: {e}")

if "report" in st.session_state:
  st.markdown("### 📋 AI 점검 리포트")
  st.markdown(st.session_state["report"])

  st.divider()
  st.subheader("📥 점검 리포트 다운로드")

  col1, col2 = st.columns(2)

  with col1:
    # Word 다운로드 버튼
    docx_bytes = create_docx(st.session_state["report"])
    st.download_button(
        label="📄 Word 파일 (.docx) 다운로드",
        data=docx_bytes,
        file_name="간호교육_포트폴리오_점검리포트.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )

  with col2:
    # TXT 다운로드 버튼
    st.download_button(
        label="📝 텍스트 파일 (.txt) 다운로드",
        data=st.session_state["report"],
        file_name="간호교육_포트폴리오_점검리포트.txt",
        mime="text/plain",
    )
